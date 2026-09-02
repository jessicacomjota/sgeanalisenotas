import io
import re
import hmac
from pathlib import Path

import pandas as pd
import streamlit as st

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak
)


# ============================================================
# CONFIGURAÇÃO
# ============================================================

st.set_page_config(
    page_title="Sistema de Análise de Notas - SESI",
    page_icon="📚",
    layout="wide"
)

CAMINHO_LOGO = Path(__file__).parent / "logoSesi.jpg"

MEDIA_MINIMA = 70
META_RECUPERACAO = 140

BIMESTRES = [
    "1º Bimestre",
    "2º Bimestre",
    "3º Bimestre",
    "4º Bimestre"
]

SITUACOES_RECUPERACAO = [
    "RECUPERAÇÃO",
    "CONTINUA ABAIXO DE 70"
]

SITUACOES_FINAL = [
    "RECUPERAÇÃO FINAL",
    "REPROVADO APÓS RECUPERAÇÃO FINAL"
]

PERFIL_ADMINISTRADOR = "Administrador"
PERFIL_COORDENACAO = "Coordenação"


# ============================================================
# AUTENTICAÇÃO
# ============================================================

def carregar_usuarios():
    """Carrega usuários e senhas a partir do Streamlit Secrets."""

    usuarios = {}

    try:
        dados_usuarios = st.secrets["usuarios"]

        for usuario, dados in dados_usuarios.items():
            usuarios[usuario] = {
                "senha": str(dados["senha"]),
                "perfil": str(dados["perfil"])
            }

    except Exception:
        st.error(
            "Não foi possível carregar os usuários do sistema. "
            "Verifique o arquivo .streamlit/secrets.toml."
        )
        st.stop()

    return usuarios


def inicializar_sessao():

    if "usuarios_app" not in st.session_state:
        st.session_state.usuarios_app = carregar_usuarios()

    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False

    if "usuario_logado" not in st.session_state:
        st.session_state.usuario_logado = None

    if "perfil_logado" not in st.session_state:
        st.session_state.perfil_logado = None

    if "bimestre_selecionado" not in st.session_state:
        st.session_state.bimestre_selecionado = "1º Bimestre"


def autenticar_usuario(usuario, senha):

    usuario = usuario.strip().lower()
    usuarios = st.session_state.usuarios_app

    if usuario not in usuarios:
        return False

    senha_cadastrada = str(
        usuarios[usuario].get("senha", "")
    )

    return hmac.compare_digest(
        str(senha),
        senha_cadastrada
    )


def realizar_login(usuario, senha):

    usuario = usuario.strip().lower()

    if autenticar_usuario(usuario, senha):

        dados_usuario = st.session_state.usuarios_app[usuario]

        st.session_state.autenticado = True
        st.session_state.usuario_logado = usuario
        st.session_state.perfil_logado = dados_usuario["perfil"]

        return True

    return False


def realizar_logout():

    st.session_state.autenticado = False
    st.session_state.usuario_logado = None
    st.session_state.perfil_logado = None

    chaves = [
        "bimestre_selecionado",
        "periodo_relatorio",
        "bimestre_relatorio"
    ]

    for chave in chaves:
        if chave in st.session_state:
            del st.session_state[chave]

    st.rerun()


def usuario_eh_administrador():

    return (
        st.session_state.get("perfil_logado")
        == PERFIL_ADMINISTRADOR
    )


# ============================================================
# TELA DE LOGIN
# ============================================================

def mostrar_tela_login():

    st.markdown(
        """
        <style>

        .login-titulo {
            color: #005DAA;
            text-align: center;
            font-size: 2rem;
            font-weight: 700;
            margin-bottom: 5px;
        }

        .login-subtitulo {
            text-align: center;
            color: #555555;
            margin-bottom: 25px;
        }

        </style>
        """,
        unsafe_allow_html=True
    )

    _, col_centro, _ = st.columns([1, 1.5, 1])

    with col_centro:

        if CAMINHO_LOGO.exists():

            _, col_logo, _ = st.columns([1, 2, 1])

            with col_logo:

                st.image(
                    str(CAMINHO_LOGO),
                    use_container_width=True
                )

        st.markdown(
            """
            <div class="login-titulo">
                Sistema de Análise de Notas
            </div>

            <div class="login-subtitulo">
                SESI Escola
            </div>
            """,
            unsafe_allow_html=True
        )

        with st.form(
            "formulario_login",
            clear_on_submit=False
        ):

            usuario = st.text_input(
                "Usuário",
                placeholder="Digite seu usuário"
            )

            senha = st.text_input(
                "Senha",
                type="password",
                placeholder="Digite sua senha"
            )

            entrar = st.form_submit_button(
                "🔐 Entrar",
                use_container_width=True,
                type="primary"
            )

        if entrar:

            if not usuario.strip():

                st.warning("Informe o usuário.")

            elif not senha:

                st.warning("Informe a senha.")

            elif realizar_login(usuario, senha):

                st.success(
                    "Login realizado com sucesso."
                )

                st.rerun()

            else:

                st.error(
                    "Usuário ou senha incorretos."
                )


# ============================================================
# GERENCIAMENTO DE USUÁRIOS
# ============================================================

def gerenciar_usuarios():

    if not usuario_eh_administrador():
        return

    st.sidebar.divider()

    with st.sidebar.expander(
        "👥 Gerenciar usuários",
        expanded=False
    ):

        st.caption(
            "Área exclusiva do administrador."
        )

        st.markdown(
            "**Usuários cadastrados**"
        )

        usuarios_ordenados = sorted(
            st.session_state.usuarios_app.keys()
        )

        for usuario in usuarios_ordenados:

            dados = (
                st.session_state
                .usuarios_app[usuario]
            )

            st.write(
                f"• **{usuario}** — "
                f"{dados['perfil']}"
            )

        st.divider()

        st.markdown(
            "**Adicionar usuário**"
        )

        novo_usuario = st.text_input(
            "Novo usuário",
            key="admin_novo_usuario"
        )

        nova_senha = st.text_input(
            "Senha",
            type="password",
            key="admin_nova_senha"
        )

        novo_perfil = st.selectbox(
            "Perfil",
            [
                PERFIL_COORDENACAO,
                PERFIL_ADMINISTRADOR
            ],
            key="admin_novo_perfil"
        )

        if st.button(
            "➕ Adicionar usuário",
            key="botao_adicionar_usuario",
            use_container_width=True
        ):

            usuario_normalizado = (
                novo_usuario.strip().lower()
            )

            if not usuario_normalizado:

                st.warning(
                    "Informe o nome do usuário."
                )

            elif not nova_senha:

                st.warning(
                    "Informe uma senha."
                )

            elif usuario_normalizado in (
                st.session_state.usuarios_app
            ):

                st.error(
                    "Esse usuário já existe."
                )

            else:

                st.session_state.usuarios_app[
                    usuario_normalizado
                ] = {
                    "senha": nova_senha,
                    "perfil": novo_perfil
                }

                st.success(
                    f"Usuário '{usuario_normalizado}' "
                    "adicionado com sucesso."
                )

                st.rerun()

        st.divider()

        st.markdown(
            "**Remover usuário**"
        )

        usuarios_removiveis = [
            usuario
            for usuario in
            st.session_state.usuarios_app.keys()
            if (
                usuario
                != st.session_state.usuario_logado
                and usuario != "jessicamartins"
            )
        ]

        if usuarios_removiveis:

            usuario_remover = st.selectbox(
                "Selecione o usuário",
                sorted(usuarios_removiveis),
                key="admin_usuario_remover"
            )

            if st.button(
                "🗑️ Remover usuário",
                key="botao_remover_usuario",
                use_container_width=True
            ):

                del st.session_state.usuarios_app[
                    usuario_remover
                ]

                st.success(
                    f"Usuário '{usuario_remover}' removido."
                )

                st.rerun()

        else:

            st.info(
                "Não há usuários disponíveis para remoção."
            )


# ============================================================
# USUÁRIO LOGADO
# ============================================================

def mostrar_usuario_logado():

    usuario = st.session_state.usuario_logado
    perfil = st.session_state.perfil_logado

    st.sidebar.markdown("### 👤 Usuário")

    st.sidebar.write(
        f"**{usuario}**"
    )

    st.sidebar.caption(
        f"Perfil: {perfil}"
    )

    if usuario_eh_administrador():

        st.sidebar.success(
            "🔑 Acesso de administrador"
        )

    else:

        st.sidebar.info(
            "👤 Acesso da coordenação"
        )

    if st.sidebar.button(
        "🚪 Sair",
        key="botao_logout",
        use_container_width=True
    ):

        realizar_logout()

    gerenciar_usuarios()


# ============================================================
# INICIAR AUTENTICAÇÃO
# ============================================================

inicializar_sessao()

if not st.session_state.autenticado:

    mostrar_tela_login()
    st.stop()

mostrar_usuario_logado()


# ============================================================
# ESTILO VISUAL
# ============================================================

st.markdown(
    """
    <style>

    h1 {
        color: #005DAA !important;
        font-weight: 700 !important;
    }

    .titulo-secao {
        background-color: #005DAA !important;
        color: #FFFFFF !important;
        font-size: 1.35rem !important;
        font-weight: 700 !important;
        padding: 12px 18px !important;
        margin-top: 25px !important;
        margin-bottom: 15px !important;
        line-height: 1.3 !important;
        border-radius: 6px !important;
        width: 100% !important;
        box-sizing: border-box !important;
        display: block !important;
    }

    .subtitulo-secao {
        background-color: #005DAA !important;
        color: #FFFFFF !important;
        font-size: 1.10rem !important;
        font-weight: 700 !important;
        padding: 9px 15px !important;
        margin-top: 15px !important;
        margin-bottom: 10px !important;
        line-height: 1.3 !important;
        border-radius: 5px !important;
        width: 100% !important;
        box-sizing: border-box !important;
        display: block !important;
    }

    .stCaption {
        font-size: 0.95rem !important;
    }

    div.stButton > button {
        border-radius: 8px !important;
        font-weight: 600 !important;
    }

    hr {
        margin-top: 25px !important;
        margin-bottom: 25px !important;
    }

    </style>
    """,
    unsafe_allow_html=True
)


# ============================================================
# CABEÇALHO
# ============================================================

col_logo, col_titulo = st.columns([1, 5])

with col_logo:

    if CAMINHO_LOGO.exists():

        st.image(
            str(CAMINHO_LOGO),
            width=160
        )

with col_titulo:

    st.title(
        "Sistema de Análise de Notas"
    )

    st.caption(
        "SESI Escola — Análise acadêmica "
        "de notas, recuperação e desempenho"
    )

st.divider()


# ============================================================
# UPLOAD
# ============================================================

st.write(
    """
    Envie o relatório de notas do SGE em formato CSV.

    O sistema realiza automaticamente a análise dos
    bimestres, recuperações, médias semestrais,
    média anual e situação final dos alunos.
    """
)

arquivo = st.file_uploader(
    "📂 Selecione o relatório CSV do SGE",
    type=["csv"]
)


# ============================================================
# FUNÇÕES AUXILIARES
# ============================================================

def limpar(valor):

    if pd.isna(valor):
        return ""

    valor = str(valor)

    valor = valor.replace(
        "\xa0",
        " "
    )

    valor = valor.replace(
        "\ufeff",
        ""
    )

    valor = re.sub(
        r"\s+",
        " ",
        valor
    )

    return valor.strip()


def converter_nota(valor):

    valor = limpar(valor)

    if valor in [
        "",
        "--",
        "----",
        "-",
        "nan",
        "None"
    ]:

        return None

    valor = valor.replace(
        ",",
        "."
    )

    try:

        return float(valor)

    except (
        ValueError,
        TypeError
    ):

        return None


def formatar_nota(valor):

    if valor is None or pd.isna(valor):
        return ""

    try:

        return (
            f"{float(valor):.2f}"
            .replace(".", ",")
        )

    except Exception:

        return ""


def nota_existe(valor):

    return (
        valor is not None
        and not pd.isna(valor)
    )


# ============================================================
# LEITURA CSV
# ============================================================

def ler_csv(uploaded_file):

    conteudo = uploaded_file.getvalue()

    codificacoes = [
        "utf-8-sig",
        "cp1252",
        "latin1"
    ]

    ultimo_erro = None

    for codificacao in codificacoes:

        try:

            df = pd.read_csv(
                io.BytesIO(conteudo),
                sep=";",
                encoding=codificacao,
                header=None,
                dtype=str
            )

            return (
                df.fillna(""),
                codificacao
            )

        except UnicodeDecodeError as erro:

            ultimo_erro = erro

    raise ValueError(
        "Não foi possível identificar "
        "a codificação do arquivo."
    ) from ultimo_erro


# ============================================================
# IDENTIFICAR DISCIPLINAS
# ============================================================

def identificar_disciplinas(df):

    disciplinas = {}

    limite = min(
        15,
        len(df)
    )

    for coluna in range(
        len(df.columns)
    ):

        for linha in range(limite):

            texto = limpar(
                df.iloc[linha, coluna]
            )

            match = re.match(
                r"^(\d{2,3})\s*-\s*(.+)$",
                texto
            )

            if not match:
                continue

            nome = match.group(2).strip()

            if nome.upper() in [
                "MAT-SGA",
                "VES-SGA",
                "MAP"
            ]:
                continue

            disciplinas[coluna] = nome

            break

    return disciplinas


# ============================================================
# LOCALIZAR ETAPA
# ============================================================

def localizar_coluna_etapa(df):

    etapas = [
        "1º Bimestre",
        "2º Bimestre",
        "3º Bimestre",
        "4º Bimestre",
        "Nota Recuperação - 1º Semestre",
        "Nota Recuperação - 2º Semestre",
        "Recuperação Final"
    ]

    if len(df.columns) > 6:

        valores = [
            limpar(x)
            for x in df.iloc[:, 6].tolist()
        ]

        encontrados = sum(
            1
            for etapa in etapas
            if etapa in valores
        )

        if encontrados >= 2:
            return 6

    for coluna in range(
        len(df.columns)
    ):

        valores = [
            limpar(x)
            for x in df.iloc[:, coluna].tolist()
        ]

        encontrados = sum(
            1
            for etapa in etapas
            if etapa in valores
        )

        if encontrados >= 2:
            return coluna

    return None


# ============================================================
# ENCONTRAR ETAPA
# ============================================================

def encontrar_linha_etapa(
    df,
    inicio,
    fim,
    coluna_etapa,
    etapa_procurada
):

    fim = min(
        fim,
        len(df)
    )

    for j in range(
        inicio,
        fim
    ):

        etapa = limpar(
            df.iloc[j, coluna_etapa]
        )

        if etapa == etapa_procurada:

            return j

    return None


# ============================================================
# PROCESSAR NOTAS
# ============================================================

def processar_notas(
    df,
    disciplinas,
    coluna_etapa
):

    registros = []

    aluno_atual = None

    i = 0

    while i < len(df):

        primeira_coluna = limpar(
            df.iloc[i, 0]
        )

        if primeira_coluna:

            texto_lower = (
                primeira_coluna.lower()
            )

            if (
                primeira_coluna != "Nome"
                and "relatório" not in texto_lower
                and "curso" not in texto_lower
                and "turno" not in texto_lower
                and "série" not in texto_lower
                and "turma" not in texto_lower
                and "aluno" not in texto_lower
            ):

                aluno_atual = primeira_coluna

        etapa = limpar(
            df.iloc[i, coluna_etapa]
        )

        # ====================================================
        # PRIMEIRO SEMESTRE
        # ====================================================

        if (
            aluno_atual
            and etapa == "1º Bimestre"
        ):

            linha_1 = i

            linha_2 = encontrar_linha_etapa(
                df,
                i + 1,
                i + 6,
                coluna_etapa,
                "2º Bimestre"
            )

            if linha_2 is not None:

                linha_rec1 = encontrar_linha_etapa(
                    df,
                    linha_2 + 1,
                    linha_2 + 8,
                    coluna_etapa,
                    "Nota Recuperação - 1º Semestre"
                )

                for coluna, disciplina in disciplinas.items():

                    nota_1 = converter_nota(
                        df.iloc[
                            linha_1,
                            coluna
                        ]
                    )

                    nota_2 = converter_nota(
                        df.iloc[
                            linha_2,
                            coluna
                        ]
                    )

                    nota_rec1 = None

                    if linha_rec1 is not None:

                        nota_rec1 = converter_nota(
                            df.iloc[
                                linha_rec1,
                                coluna
                            ]
                        )

                    if (
                        nota_1 is not None
                        and nota_2 is not None
                    ):

                        media_semestre1 = round(
                            (
                                nota_1
                                + nota_2
                            ) / 2,
                            2
                        )

                    else:

                        media_semestre1 = None

                    if (
                        media_semestre1 is not None
                        and media_semestre1 < MEDIA_MINIMA
                    ):

                        nota_necessaria_rec1 = round(
                            META_RECUPERACAO
                            - media_semestre1,
                            2
                        )

                        situacao_antes_rec1 = (
                            "RECUPERAÇÃO"
                        )

                    elif media_semestre1 is not None:

                        nota_necessaria_rec1 = None

                        situacao_antes_rec1 = (
                            "APROVADO"
                        )

                    else:

                        nota_necessaria_rec1 = None

                        situacao_antes_rec1 = (
                            "AGUARDANDO NOTAS"
                        )

                    if (
                        nota_rec1 is not None
                        and media_semestre1 is not None
                    ):

                        media_pos_rec1 = round(
                            (
                                media_semestre1
                                + nota_rec1
                            ) / 2,
                            2
                        )

                        if (
                            media_pos_rec1
                            >= MEDIA_MINIMA
                        ):

                            situacao_pos_rec1 = (
                                "APROVADO APÓS RECUPERAÇÃO"
                            )

                        else:

                            situacao_pos_rec1 = (
                                "CONTINUA ABAIXO DE 70"
                            )

                    else:

                        media_pos_rec1 = None

                        situacao_pos_rec1 = (
                            situacao_antes_rec1
                        )

                    registros.append({

                        "Aluno":
                            aluno_atual,

                        "Disciplina":
                            disciplina,

                        "1º Bimestre":
                            nota_1,

                        "2º Bimestre":
                            nota_2,

                        "Média 1º Semestre":
                            media_semestre1,

                        "Recuperação 1º Semestre":
                            nota_rec1,

                        "Média após Recuperação 1º Semestre":
                            media_pos_rec1,

                        "Nota necessária Recuperação 1º Semestre":
                            nota_necessaria_rec1,

                        "Situação 1º Semestre":
                            situacao_pos_rec1
                    })


        # ====================================================
        # SEGUNDO SEMESTRE
        # ====================================================

        elif (
            aluno_atual
            and etapa == "3º Bimestre"
        ):

            linha_3 = i

            linha_4 = encontrar_linha_etapa(
                df,
                i + 1,
                i + 6,
                coluna_etapa,
                "4º Bimestre"
            )

            if linha_4 is not None:

                linha_rec2 = encontrar_linha_etapa(
                    df,
                    linha_4 + 1,
                    linha_4 + 8,
                    coluna_etapa,
                    "Nota Recuperação - 2º Semestre"
                )

                linha_final = encontrar_linha_etapa(
                    df,
                    linha_4 + 1,
                    min(
                        linha_4 + 12,
                        len(df)
                    ),
                    coluna_etapa,
                    "Recuperação Final"
                )

                for coluna, disciplina in disciplinas.items():

                    nota_3 = converter_nota(
                        df.iloc[
                            linha_3,
                            coluna
                        ]
                    )

                    nota_4 = converter_nota(
                        df.iloc[
                            linha_4,
                            coluna
                        ]
                    )

                    nota_rec2 = None

                    if linha_rec2 is not None:

                        nota_rec2 = converter_nota(
                            df.iloc[
                                linha_rec2,
                                coluna
                            ]
                        )

                    nota_final = None

                    if linha_final is not None:

                        nota_final = converter_nota(
                            df.iloc[
                                linha_final,
                                coluna
                            ]
                        )

                    if (
                        nota_3 is not None
                        and nota_4 is not None
                    ):

                        media_semestre2 = round(
                            (
                                nota_3
                                + nota_4
                            ) / 2,
                            2
                        )

                    else:

                        media_semestre2 = None

                    if (
                        media_semestre2 is not None
                        and media_semestre2 < MEDIA_MINIMA
                    ):

                        nota_necessaria_rec2 = round(
                            META_RECUPERACAO
                            - media_semestre2,
                            2
                        )

                        situacao_antes_rec2 = (
                            "RECUPERAÇÃO"
                        )

                    elif media_semestre2 is not None:

                        nota_necessaria_rec2 = None

                        situacao_antes_rec2 = (
                            "APROVADO"
                        )

                    else:

                        nota_necessaria_rec2 = None

                        situacao_antes_rec2 = (
                            "AGUARDANDO NOTAS"
                        )

                    if (
                        nota_rec2 is not None
                        and media_semestre2 is not None
                    ):

                        media_pos_rec2 = round(
                            (
                                media_semestre2
                                + nota_rec2
                            ) / 2,
                            2
                        )

                        if (
                            media_pos_rec2
                            >= MEDIA_MINIMA
                        ):

                            situacao_pos_rec2 = (
                                "APROVADO APÓS RECUPERAÇÃO"
                            )

                        else:

                            situacao_pos_rec2 = (
                                "CONTINUA ABAIXO DE 70"
                            )

                    else:

                        media_pos_rec2 = None

                        situacao_pos_rec2 = (
                            situacao_antes_rec2
                        )

                    registros_primeiro = [

                        r
                        for r in registros

                        if (
                            r["Aluno"]
                            == aluno_atual
                            and r["Disciplina"]
                            == disciplina
                        )
                    ]

                    media_1 = None

                    if registros_primeiro:

                        registro_1 = (
                            registros_primeiro[-1]
                        )

                        media_1 = registro_1[
                            "Média após Recuperação 1º Semestre"
                        ]

                        if media_1 is None:

                            media_1 = registro_1[
                                "Média 1º Semestre"
                            ]

                    media_2 = media_pos_rec2

                    if media_2 is None:

                        media_2 = media_semestre2

                    if (
                        media_1 is not None
                        and media_2 is not None
                    ):

                        media_anual = round(
                            (
                                media_1
                                + media_2
                            ) / 2,
                            2
                        )

                    else:

                        media_anual = None

                    if media_anual is None:

                        situacao_anual = (
                            "AGUARDANDO NOTAS"
                        )

                    elif (
                        media_anual
                        >= MEDIA_MINIMA
                    ):

                        situacao_anual = (
                            "APROVADO"
                        )

                    else:

                        situacao_anual = (
                            "RECUPERAÇÃO FINAL"
                        )

                    if (
                        media_anual is not None
                        and media_anual < MEDIA_MINIMA
                    ):

                        nota_necessaria_final = round(
                            META_RECUPERACAO
                            - media_anual,
                            2
                        )

                    else:

                        nota_necessaria_final = None

                    if (
                        nota_final is not None
                        and media_anual is not None
                    ):

                        media_pos_final = round(
                            (
                                media_anual
                                + nota_final
                            ) / 2,
                            2
                        )

                        if (
                            media_pos_final
                            >= MEDIA_MINIMA
                        ):

                            situacao_final = (
                                "APROVADO APÓS RECUPERAÇÃO FINAL"
                            )

                        else:

                            situacao_final = (
                                "REPROVADO APÓS RECUPERAÇÃO FINAL"
                            )

                    else:

                        media_pos_final = None

                        if media_anual is None:

                            situacao_final = (
                                "AGUARDANDO MÉDIA ANUAL"
                            )

                        else:

                            situacao_final = (
                                situacao_anual
                            )

                    registros.append({

                        "Aluno":
                            aluno_atual,

                        "Disciplina":
                            disciplina,

                        "3º Bimestre":
                            nota_3,

                        "4º Bimestre":
                            nota_4,

                        "Média 2º Semestre":
                            media_semestre2,

                        "Recuperação 2º Semestre":
                            nota_rec2,

                        "Média após Recuperação 2º Semestre":
                            media_pos_rec2,

                        "Nota necessária Recuperação 2º Semestre":
                            nota_necessaria_rec2,

                        "Média Anual":
                            media_anual,

                        "Recuperação Final":
                            nota_final,

                        "Nota necessária Recuperação Final":
                            nota_necessaria_final,

                        "Média após Recuperação Final":
                            media_pos_final,

                        "Situação 2º Semestre":
                            situacao_pos_rec2,

                        "Situação Final":
                            situacao_final
                    })

        i += 1

    return pd.DataFrame(registros)


# ============================================================
# GARANTIR COLUNAS
# ============================================================

def garantir_colunas(analise):

    colunas_obrigatorias = [

        "Aluno",
        "Disciplina",

        "1º Bimestre",
        "2º Bimestre",

        "Média 1º Semestre",
        "Recuperação 1º Semestre",

        "Média após Recuperação 1º Semestre",

        "Nota necessária Recuperação 1º Semestre",

        "Situação 1º Semestre",

        "3º Bimestre",
        "4º Bimestre",

        "Média 2º Semestre",
        "Recuperação 2º Semestre",

        "Média após Recuperação 2º Semestre",

        "Nota necessária Recuperação 2º Semestre",

        "Média Anual",

        "Recuperação Final",

        "Nota necessária Recuperação Final",

        "Média após Recuperação Final",

        "Situação 2º Semestre",

        "Situação Final"
    ]

    for coluna in colunas_obrigatorias:

        if coluna not in analise.columns:

            analise[coluna] = None

    return analise


# ============================================================
# RANKING
# ============================================================

def gerar_ranking_bimestre(
    analise,
    bimestre,
    disciplinas
):

    if (
        analise.empty
        or bimestre not in analise.columns
    ):

        return pd.DataFrame()

    total_notas = (
        analise[bimestre]
        .notna()
        .sum()
    )

    if total_notas == 0:

        return pd.DataFrame()

    base = analise[
        [
            "Aluno",
            "Disciplina",
            bimestre
        ]
    ].copy()

    total_disciplinas = len(
        disciplinas
    )

    if total_disciplinas == 0:

        return pd.DataFrame()

    base[bimestre] = pd.to_numeric(
        base[bimestre],
        errors="coerce"
    )

    ranking = (
        base
        .groupby("Aluno")
        .agg(
            Notas_Lançadas=(
                bimestre,
                lambda x: x.notna().sum()
            ),

            Soma_Notas=(
                bimestre,
                "sum"
            )
        )
        .reset_index()
    )

    ranking = ranking[
        ranking["Notas_Lançadas"]
        >= total_disciplinas
    ].copy()

    if ranking.empty:

        return pd.DataFrame()

    ranking[
        "Média do Bimestre"
    ] = (
        ranking["Soma_Notas"]
        / ranking["Notas_Lançadas"]
    ).round(2)

    ranking = (
        ranking
        .sort_values(
            "Média do Bimestre",
            ascending=False
        )
        .head(5)
        .reset_index(drop=True)
    )

    ranking.insert(
        0,
        "Posição",
        range(
            1,
            len(ranking) + 1
        )
    )

    ranking = ranking[
        [
            "Posição",
            "Aluno",
            "Média do Bimestre",
            "Notas_Lançadas"
        ]
    ]

    ranking.rename(
        columns={
            "Notas_Lançadas":
                "Disciplinas com Nota"
        },
        inplace=True
    )

    return ranking


# ============================================================
# RESUMO DOS ALUNOS
# ============================================================

def gerar_resumo_alunos(analise):

    if analise.empty:

        return pd.DataFrame()

    def contar_situacoes(
        serie,
        situacoes
    ):

        return serie.isin(
            situacoes
        ).sum()

    agregacoes = {

        "Disciplinas":
            (
                "Disciplina",
                "nunique"
            )
    }

    if (
        "Situação 1º Semestre"
        in analise.columns
    ):

        agregacoes[
            "Recuperações_1º_Semestre"
        ] = (
            "Situação 1º Semestre",
            lambda x:
                contar_situacoes(
                    x,
                    SITUACOES_RECUPERACAO
                )
        )

    if (
        "Situação 2º Semestre"
        in analise.columns
    ):

        agregacoes[
            "Recuperações_2º_Semestre"
        ] = (
            "Situação 2º Semestre",
            lambda x:
                contar_situacoes(
                    x,
                    SITUACOES_RECUPERACAO
                )
        )

    if (
        "Situação Final"
        in analise.columns
    ):

        agregacoes[
            "Recuperações_Finais"
        ] = (
            "Situação Final",
            lambda x:
                contar_situacoes(
                    x,
                    SITUACOES_FINAL
                )
        )

    return (
        analise
        .groupby("Aluno")
        .agg(**agregacoes)
        .reset_index()
    )


# ============================================================
# RESUMO DAS DISCIPLINAS
# ============================================================

def gerar_resumo_disciplinas(
    analise
):

    if analise.empty:

        return pd.DataFrame()

    agregacoes = {

        "Alunos":
            (
                "Aluno",
                "nunique"
            )
    }

    if (
        "Situação 1º Semestre"
        in analise.columns
    ):

        agregacoes[
            "Recuperação_1º_Semestre"
        ] = (
            "Situação 1º Semestre",
            lambda x:
                x.isin(
                    SITUACOES_RECUPERACAO
                ).sum()
        )

    if (
        "Situação 2º Semestre"
        in analise.columns
    ):

        agregacoes[
            "Recuperação_2º_Semestre"
        ] = (
            "Situação 2º Semestre",
            lambda x:
                x.isin(
                    SITUACOES_RECUPERACAO
                ).sum()
        )

    if (
        "Situação Final"
        in analise.columns
    ):

        agregacoes[
            "Recuperação_Final"
        ] = (
            "Situação Final",
            lambda x:
                x.isin(
                    SITUACOES_FINAL
                ).sum()
        )

    return (
        analise
        .groupby("Disciplina")
        .agg(**agregacoes)
        .reset_index()
    )


# ============================================================
# FILTROS
# ============================================================

def aplicar_filtros(
    df,
    busca_aluno="",
    disciplinas=None,
    status=None
):

    resultado = df.copy()

    if busca_aluno:

        resultado = resultado[
            resultado["Aluno"]
            .astype(str)
            .str.contains(
                busca_aluno,
                case=False,
                na=False
            )
        ]

    if disciplinas:

        resultado = resultado[
            resultado["Disciplina"]
            .isin(disciplinas)
        ]

    if status:

        colunas_status = [

            coluna

            for coluna in [

                "Situação 1º Semestre",
                "Situação 2º Semestre",
                "Situação Final"

            ]

            if coluna in resultado.columns
        ]

        if colunas_status:

            mascara = pd.Series(
                False,
                index=resultado.index
            )

            for coluna in colunas_status:

                mascara = (
                    mascara
                    | resultado[coluna]
                    .isin(status)
                )

            resultado = resultado[
                mascara
            ]

    return resultado


def mostrar_filtros(
    analise,
    chave
):

    st.markdown(
        '<div class="subtitulo-secao">'
        'Filtros de pesquisa'
        '</div>',
        unsafe_allow_html=True
    )

    col1, col2, col3 = st.columns(
        [2, 1.5, 1.5]
    )

    with col1:

        busca = st.text_input(
            "Nome do aluno",
            placeholder=
            "Digite o nome do aluno...",
            key=f"busca_aluno_{chave}"
        )

    with col2:

        opcoes_disciplinas = sorted(
            analise[
                "Disciplina"
            ]
            .dropna()
            .unique()
            .tolist()
        )

        disciplinas = st.multiselect(
            "Disciplina",
            options=opcoes_disciplinas,
            key=f"disciplinas_{chave}"
        )

    with col3:

        opcoes_status = set()

        for coluna in [

            "Situação 1º Semestre",
            "Situação 2º Semestre",
            "Situação Final"

        ]:

            if coluna in analise.columns:

                opcoes_status.update(
                    analise[coluna]
                    .dropna()
                    .astype(str)
                    .tolist()
                )

        status = st.multiselect(
            "Situação",
            options=sorted(opcoes_status),
            key=f"status_{chave}"
        )

    resultado = aplicar_filtros(
        analise,
        busca,
        disciplinas,
        status
    )

    st.caption(
        f"{len(resultado)} registros encontrados."
    )

    return resultado


# ============================================================
# RECUPERAÇÃO
# ============================================================

def preparar_relatorio_recuperacao(
    analise,
    semestre=1
):

    if semestre == 1:

        coluna_situacao = (
            "Situação 1º Semestre"
        )

        coluna_media = (
            "Média 1º Semestre"
        )

        coluna_recuperacao = (
            "Recuperação 1º Semestre"
        )

        coluna_necessaria = (
            "Nota necessária "
            "Recuperação 1º Semestre"
        )

        coluna_pos = (
            "Média após Recuperação "
            "1º Semestre"
        )

    else:

        coluna_situacao = (
            "Situação 2º Semestre"
        )

        coluna_media = (
            "Média 2º Semestre"
        )

        coluna_recuperacao = (
            "Recuperação 2º Semestre"
        )

        coluna_necessaria = (
            "Nota necessária "
            "Recuperação 2º Semestre"
        )

        coluna_pos = (
            "Média após Recuperação "
            "2º Semestre"
        )

    if (
        coluna_situacao
        not in analise.columns
    ):

        return pd.DataFrame()

    resultado = analise[
        analise[coluna_situacao]
        .isin(
            SITUACOES_RECUPERACAO
        )
    ].copy()

    if resultado.empty:

        return pd.DataFrame()

    resultado["Prova"] = (
        resultado[coluna_recuperacao]
        .apply(
            lambda x:
                "PROVA REALIZADA"
                if nota_existe(x)
                else "PROVA NÃO REALIZADA"
        )
    )

    resultado["Resultado"] = (
        resultado.apply(

            lambda linha:

                (
                    "AGUARDANDO RECUPERAÇÃO"

                    if not nota_existe(
                        linha[coluna_recuperacao]
                    )

                    else (

                        "APROVADO APÓS RECUPERAÇÃO"

                        if (
                            nota_existe(
                                linha[coluna_pos]
                            )

                            and linha[coluna_pos]
                            >= MEDIA_MINIMA
                        )

                        else
                        "CONTINUA ABAIXO DE 70"
                    )
                ),

            axis=1
        )
    )

    resultado = resultado[

        [
            "Aluno",
            "Disciplina",
            coluna_media,
            coluna_necessaria,
            coluna_recuperacao,
            "Prova",
            coluna_pos,
            "Resultado"
        ]

    ].copy()

    resultado.rename(

        columns={

            coluna_media:
                "Média Semestral",

            coluna_necessaria:
                "Nota Necessária",

            coluna_recuperacao:
                "Nota Recuperação",

            coluna_pos:
                "Média Após Recuperação"
        },

        inplace=True
    )

    return resultado


def gerar_recuperacao_por_disciplina(
    analise,
    semestre=1
):

    coluna_situacao = (

        "Situação 1º Semestre"

        if semestre == 1

        else
        "Situação 2º Semestre"
    )

    if (
        coluna_situacao
        not in analise.columns
    ):

        return pd.DataFrame()

    recuperacao = analise[
        analise[coluna_situacao]
        .isin(
            SITUACOES_RECUPERACAO
        )
    ].copy()

    if recuperacao.empty:

        return pd.DataFrame()

    resultado = (
        recuperacao
        .groupby("Disciplina")
        .agg(
            Alunos=(
                "Aluno",
                "nunique"
            )
        )
        .reset_index()
        .sort_values(
            "Alunos",
            ascending=False
        )
    )

    return resultado


def gerar_lista_alunos_por_disciplina(
    analise,
    semestre=1
):

    coluna_situacao = (

        "Situação 1º Semestre"

        if semestre == 1

        else
        "Situação 2º Semestre"
    )

    if (
        coluna_situacao
        not in analise.columns
    ):

        return pd.DataFrame()

    recuperacao = analise[
        analise[coluna_situacao]
        .isin(
            SITUACOES_RECUPERACAO
        )
    ].copy()

    if recuperacao.empty:

        return pd.DataFrame()

    def numerar_alunos(nomes):

        alunos = sorted(
            set(
                nome.strip()
                for nome in nomes
                if nome
                and str(nome).strip()
            )
        )

        return "\n".join(
            f"{numero}. {nome}"
            for numero, nome
            in enumerate(
                alunos,
                start=1
            )
        )

    resultado = (
        recuperacao
        .groupby("Disciplina")["Aluno"]
        .apply(numerar_alunos)
        .reset_index()
    )

    resultado.rename(
        columns={
            "Aluno":
                "Alunos em Recuperação"
        },
        inplace=True
    )

    return resultado


def preparar_recuperacao_final(
    analise
):

    if (
        "Situação Final"
        not in analise.columns
    ):

        return pd.DataFrame()

    resultado = analise[
        analise["Situação Final"]
        .isin(SITUACOES_FINAL)
    ].copy()

    if resultado.empty:

        return pd.DataFrame()

    resultado["Prova"] = (
        resultado["Recuperação Final"]
        .apply(
            lambda x:
                "PROVA REALIZADA"
                if nota_existe(x)
                else "PROVA NÃO REALIZADA"
        )
    )

    resultado["Resultado"] = (
        resultado.apply(

            lambda linha:

                (
                    "AGUARDANDO RECUPERAÇÃO FINAL"

                    if not nota_existe(
                        linha["Recuperação Final"]
                    )

                    else (

                        "APROVADO APÓS RECUPERAÇÃO FINAL"

                        if (
                            nota_existe(
                                linha[
                                    "Média após Recuperação Final"
                                ]
                            )

                            and linha[
                                "Média após Recuperação Final"
                            ] >= MEDIA_MINIMA
                        )

                        else
                        "REPROVADO APÓS RECUPERAÇÃO FINAL"
                    )
                ),

            axis=1
        )
    )

    return resultado[

        [
            "Aluno",
            "Disciplina",
            "Média Anual",
            "Nota necessária Recuperação Final",
            "Recuperação Final",
            "Prova",
            "Média após Recuperação Final",
            "Resultado"
        ]

    ].rename(

        columns={

            "Nota necessária Recuperação Final":
                "Nota Necessária",

            "Recuperação Final":
                "Nota Recuperação",

            "Média após Recuperação Final":
                "Média Após Recuperação"
        }
    )


# ============================================================
# WORD
# ============================================================

def gerar_relatorio_word(
    analise,
    ranking,
    periodo,
    bimestre
):

    documento = Document()

    section = documento.sections[0]

    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    titulo = documento.add_paragraph()

    titulo.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = titulo.add_run(
        "RELATÓRIO DE DESEMPENHO ACADÊMICO"
    )

    run.bold = True
    run.font.size = Pt(17)

    subtitulo = documento.add_paragraph()

    subtitulo.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    run = subtitulo.add_run(
        "SESI Escola"
    )

    run.bold = True
    run.font.size = Pt(12)

    periodo_texto = documento.add_paragraph()

    periodo_texto.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    periodo_texto.add_run(
        f"Período: {periodo}"
    ).font.size = Pt(10)

    documento.add_paragraph()

    # ========================================================
    # INDICADORES
    # ========================================================

    documento.add_heading(
        "1. INDICADORES GERAIS",
        level=1
    )

    total_alunos = (
        analise["Aluno"]
        .nunique()
    )

    total_disciplinas = (
        analise["Disciplina"]
        .nunique()
    )

    rec1 = preparar_relatorio_recuperacao(
        analise,
        1
    )

    rec2 = preparar_relatorio_recuperacao(
        analise,
        2
    )

    final = preparar_recuperacao_final(
        analise
    )

    tabela = documento.add_table(
        rows=1,
        cols=4
    )

    tabela.style = "Table Grid"

    tabela.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    cab = tabela.rows[0].cells

    cab[0].text = "Alunos"
    cab[1].text = "Disciplinas"
    cab[2].text = "Recuperações"
    cab[3].text = "Recuperações Finais"

    linha = tabela.add_row().cells

    linha[0].text = str(
        total_alunos
    )

    linha[1].text = str(
        total_disciplinas
    )

    linha[2].text = str(
        len(rec1)
        + len(rec2)
    )

    linha[3].text = str(
        len(final)
    )

    # ========================================================
    # RANKING
    # ========================================================

    documento.add_heading(
        f"2. ALUNOS DESTAQUE — {bimestre}",
        level=1
    )

    if ranking.empty:

        documento.add_paragraph(
            "Não foi possível gerar "
            "o ranking para este bimestre."
        )

    else:

        tabela = documento.add_table(
            rows=1,
            cols=4
        )

        tabela.style = "Table Grid"

        cabecalho = (
            tabela.rows[0].cells
        )

        cabecalho[0].text = "Posição"
        cabecalho[1].text = "Aluno"
        cabecalho[2].text = "Média"
        cabecalho[3].text = "Disciplinas"

        for _, linha_rank in (
            ranking.iterrows()
        ):

            cells = (
                tabela.add_row().cells
            )

            cells[0].text = str(
                linha_rank["Posição"]
            )

            cells[1].text = str(
                linha_rank["Aluno"]
            )

            cells[2].text = formatar_nota(
                linha_rank[
                    "Média do Bimestre"
                ]
            )

            cells[3].text = str(
                linha_rank[
                    "Disciplinas com Nota"
                ]
            )

    # ========================================================
    # RESUMO
    # ========================================================

    resumo_alunos = (
        gerar_resumo_alunos(
            analise
        )
    )

    documento.add_heading(
        "3. RESUMO DOS ALUNOS",
        level=1
    )

    if not resumo_alunos.empty:

        tabela = documento.add_table(
            rows=1,
            cols=len(
                resumo_alunos.columns
            )
        )

        tabela.style = "Table Grid"

        for i, coluna in enumerate(
            resumo_alunos.columns
        ):

            tabela.rows[0].cells[i].text = (
                str(coluna)
            )

        for _, linha_data in (
            resumo_alunos.iterrows()
        ):

            cells = (
                tabela.add_row().cells
            )

            for i, coluna in enumerate(
                resumo_alunos.columns
            ):

                cells[i].text = str(
                    linha_data[coluna]
                )

    buffer = io.BytesIO()

    documento.save(buffer)

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# PDF
# ============================================================

def criar_tabela_pdf(
    dados,
    fonte=7
):

    tabela = Table(
        dados,
        repeatRows=1,
        hAlign="CENTER"
    )

    tabela.setStyle(

        TableStyle([

            (
                "BACKGROUND",
                (0, 0),
                (-1, 0),
                colors.HexColor("#005DAA")
            ),

            (
                "TEXTCOLOR",
                (0, 0),
                (-1, 0),
                colors.white
            ),

            (
                "FONTNAME",
                (0, 0),
                (-1, 0),
                "Helvetica-Bold"
            ),

            (
                "GRID",
                (0, 0),
                (-1, -1),
                0.4,
                colors.HexColor("#E0E0E0")
            ),

            (
                "FONTSIZE",
                (0, 0),
                (-1, -1),
                fonte
            ),

            (
                "VALIGN",
                (0, 0),
                (-1, -1),
                "MIDDLE"
            ),

            (
                "LEFTPADDING",
                (0, 0),
                (-1, -1),
                4
            ),

            (
                "RIGHTPADDING",
                (0, 0),
                (-1, -1),
                4
            ),

            (
                "TOPPADDING",
                (0, 0),
                (-1, -1),
                4
            ),

            (
                "BOTTOMPADDING",
                (0, 0),
                (-1, -1),
                4
            )

        ])
    )

    return tabela


def gerar_relatorio_pdf(
    analise,
    ranking,
    periodo,
    bimestre
):

    buffer = io.BytesIO()

    documento = SimpleDocTemplate(

        buffer,

        pagesize=landscape(A4),

        rightMargin=1.0 * cm,
        leftMargin=1.0 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm
    )

    estilos = getSampleStyleSheet()

    estilo_titulo = ParagraphStyle(
        "TituloRelatorio",
        parent=estilos["Title"],
        alignment=TA_CENTER,
        fontSize=17,
        leading=20,
        spaceAfter=8
    )

    estilo_subtitulo = ParagraphStyle(
        "SubtituloRelatorio",
        parent=estilos["Normal"],
        alignment=TA_CENTER,
        fontSize=10,
        leading=13,
        spaceAfter=15
    )

    estilo_secao = ParagraphStyle(
        "SecaoRelatorio",
        parent=estilos["Heading2"],
        fontSize=13,
        leading=16,
        spaceBefore=8,
        spaceAfter=8
    )

    elementos = []

    elementos.append(
        Paragraph(
            "RELATÓRIO DE DESEMPENHO ACADÊMICO",
            estilo_titulo
        )
    )

    elementos.append(
        Paragraph(
            "SESI Escola",
            estilo_subtitulo
        )
    )

    elementos.append(
        Paragraph(
            f"Período: {periodo}",
            estilo_subtitulo
        )
    )

    # ========================================================
    # INDICADORES
    # ========================================================

    elementos.append(
        Paragraph(
            "1. INDICADORES GERAIS",
            estilo_secao
        )
    )

    total_alunos = (
        analise["Aluno"]
        .nunique()
    )

    total_disciplinas = (
        analise["Disciplina"]
        .nunique()
    )

    rec1 = preparar_relatorio_recuperacao(
        analise,
        1
    )

    rec2 = preparar_relatorio_recuperacao(
        analise,
        2
    )

    final = preparar_recuperacao_final(
        analise
    )

    dados = [

        [
            "Alunos",
            "Disciplinas",
            "Recuperações",
            "Recuperações Finais"
        ],

        [
            str(total_alunos),
            str(total_disciplinas),
            str(len(rec1) + len(rec2)),
            str(len(final))
        ]
    ]

    elementos.append(
        criar_tabela_pdf(
            dados,
            fonte=8
        )
    )

    elementos.append(
        Spacer(1, 10)
    )

    # ========================================================
    # RANKING
    # ========================================================

    elementos.append(
        Paragraph(
            f"2. ALUNOS DESTAQUE — {bimestre}",
            estilo_secao
        )
    )

    if ranking.empty:

        elementos.append(
            Paragraph(
                "Não foi possível gerar "
                "o ranking para este bimestre.",
                estilos["Normal"]
            )
        )

    else:

        dados = [

            [
                "Posição",
                "Aluno",
                "Média",
                "Disciplinas"
            ]
        ]

        for _, linha_rank in (
            ranking.iterrows()
        ):

            dados.append([

                str(
                    linha_rank["Posição"]
                ),

                str(
                    linha_rank["Aluno"]
                ),

                formatar_nota(
                    linha_rank[
                        "Média do Bimestre"
                    ]
                ),

                str(
                    linha_rank[
                        "Disciplinas com Nota"
                    ]
                )
            ])

        elementos.append(
            criar_tabela_pdf(
                dados,
                fonte=8
            )
        )

    # ========================================================
    # RESUMO DOS ALUNOS
    # ========================================================

    resumo_alunos = (
        gerar_resumo_alunos(
            analise
        )
    )

    elementos.append(
        PageBreak()
    )

    elementos.append(
        Paragraph(
            "3. RESUMO DOS ALUNOS",
            estilo_secao
        )
    )

    if not resumo_alunos.empty:

        dados = [
            list(
                resumo_alunos.columns
            )
        ]

        for _, linha in (
            resumo_alunos.iterrows()
        ):

            dados.append([
                str(
                    linha[coluna]
                )
                for coluna
                in resumo_alunos.columns
            ])

        elementos.append(
            criar_tabela_pdf(
                dados,
                fonte=6
            )
        )

    def adicionar_numero_pagina(
        canvas,
        doc
    ):

        canvas.saveState()

        canvas.setFont(
            "Helvetica",
            8
        )

        canvas.drawCentredString(
            landscape(A4)[0] / 2,
            0.5 * cm,
            f"SESI Escola — Página {doc.page}"
        )

        canvas.restoreState()

    documento.build(
        elementos,
        onFirstPage=adicionar_numero_pagina,
        onLaterPages=adicionar_numero_pagina
    )

    buffer.seek(0)

    return buffer.getvalue()


# ============================================================
# INTERFACE PRINCIPAL
# ============================================================

if arquivo is not None:

    try:

        # ====================================================
        # LEITURA
        # ====================================================

        df, codificacao = ler_csv(
            arquivo
        )

        st.success(
            "Arquivo carregado com sucesso! "
            f"Codificação: {codificacao}"
        )

        col1, col2, col3 = st.columns(3)

        col1.metric(
            "Linhas",
            len(df)
        )

        col2.metric(
            "Colunas",
            len(df.columns)
        )

        coluna_etapa = localizar_coluna_etapa(
            df
        )

        if coluna_etapa is None:

            st.error(
                "Não foi possível localizar "
                "a coluna de Etapa."
            )

            st.stop()

        disciplinas = identificar_disciplinas(
            df
        )

        col3.metric(
            "Disciplinas",
            len(disciplinas)
        )

        # ====================================================
        # DISCIPLINAS
        # ====================================================

        st.markdown(
            '<div class="titulo-secao">'
            'DISCIPLINAS IDENTIFICADAS'
            '</div>',
            unsafe_allow_html=True
        )

        tabela_disciplinas = pd.DataFrame({

            "Coluna":
                [
                    coluna + 1
                    for coluna
                    in disciplinas
                ],

            "Disciplina":
                [
                    nome
                    for nome
                    in disciplinas.values()
                ]
        })

        st.dataframe(
            tabela_disciplinas,
            use_container_width=True,
            hide_index=True
        )

        # ====================================================
        # PROCESSAMENTO
        # ====================================================

        analise = processar_notas(
            df,
            disciplinas,
            coluna_etapa
        )

        if analise.empty:

            st.warning(
                "Nenhum registro de notas "
                "foi encontrado."
            )

            st.stop()

        analise = garantir_colunas(
            analise
        )

        # ====================================================
        # RESUMOS
        # ====================================================

        resumo_alunos = (
            gerar_resumo_alunos(
                analise
            )
        )

        resumo_disciplinas = (
            gerar_resumo_disciplinas(
                analise
            )
        )

        # ====================================================
        # RESUMO GERAL
        # ====================================================

        st.markdown(
            '<div class="titulo-secao">'
            'RESUMO GERAL'
            '</div>',
            unsafe_allow_html=True
        )

        total_alunos = (
            analise["Aluno"]
            .nunique()
        )

        total_disciplinas = (
            analise["Disciplina"]
            .nunique()
        )

        total_registros = len(
            analise
        )

        total_rec1 = len(
            preparar_relatorio_recuperacao(
                analise,
                1
            )
        )

        total_rec2 = len(
            preparar_relatorio_recuperacao(
                analise,
                2
            )
        )

        total_final = len(
            preparar_recuperacao_final(
                analise
            )
        )

        c1, c2, c3, c4, c5 = st.columns(5)

        c1.metric(
            "👩‍🎓 Alunos",
            total_alunos
        )

        c2.metric(
            "📚 Disciplinas",
            total_disciplinas
        )

        c3.metric(
            "📝 Registros",
            total_registros
        )

        c4.metric(
            "🔄 Recuperações",
            total_rec1 + total_rec2
        )

        c5.metric(
            "⚠️ Recuperações Finais",
            total_final
        )

        st.divider()

        # ====================================================
        # ABAS PRINCIPAIS
        # ====================================================

        aba1, aba2, aba3, aba4, aba5, aba6 = st.tabs([

            "📊 RESUMOS",

            "🏆 RANKING",

            "📝 RECUPERAÇÕES",

            "🔎 ANÁLISE DETALHADA",

            "📄 RELATÓRIOS",

            "📥 EXPORTAR DADOS"

        ])

        # ====================================================
        # ABA 1 - RESUMOS
        # ====================================================

        with aba1:

            st.markdown(
                '<div class="titulo-secao">'
                'RESUMOS — ALUNOS E DISCIPLINAS'
                '</div>',
                unsafe_allow_html=True
            )

            resumo_tab1, resumo_tab2 = st.tabs([

                "👩‍🎓 RESUMO DOS ALUNOS",

                "📚 RESUMO DAS DISCIPLINAS"

            ])

            with resumo_tab1:

                st.markdown(
                    '<div class="subtitulo-secao">'
                    'Resumo por aluno'
                    '</div>',
                    unsafe_allow_html=True
                )

                if resumo_alunos.empty:

                    st.info(
                        "Nenhum resumo disponível."
                    )

                else:

                    st.dataframe(
                        resumo_alunos,
                        use_container_width=True,
                        hide_index=True
                    )

            with resumo_tab2:

                st.markdown(
                    '<div class="subtitulo-secao">'
                    'Resumo por disciplina'
                    '</div>',
                    unsafe_allow_html=True
                )

                if resumo_disciplinas.empty:

                    st.info(
                        "Nenhum resumo disponível."
                    )

                else:

                    st.dataframe(
                        resumo_disciplinas,
                        use_container_width=True,
                        hide_index=True
                    )

        # ====================================================
        # ABA 2 - RANKING
        # ====================================================

        with aba2:

            st.markdown(
                '<div class="titulo-secao">'
                'ALUNOS DESTAQUE — TOP 5'
                '</div>',
                unsafe_allow_html=True
            )

            b1, b2, b3, b4 = st.columns(4)

            botoes = [

                (
                    b1,
                    "1º Bimestre"
                ),

                (
                    b2,
                    "2º Bimestre"
                ),

                (
                    b3,
                    "3º Bimestre"
                ),

                (
                    b4,
                    "4º Bimestre"
                )
            ]

            for coluna, bimestre_nome in botoes:

                with coluna:

                    if st.button(

                        bimestre_nome,

                        use_container_width=True,

                        type=(
                            "primary"
                            if (
                                st.session_state
                                .bimestre_selecionado
                                == bimestre_nome
                            )
                            else
                            "secondary"
                        ),

                        key=(
                            f"ranking_{bimestre_nome}"
                        )
                    ):

                        st.session_state[
                            "bimestre_selecionado"
                        ] = bimestre_nome

            bimestre_selecionado = (
                st.session_state
                .bimestre_selecionado
            )

            st.markdown(
                f'<div class="subtitulo-secao">'
                f'RANKING — {bimestre_selecionado}'
                '</div>',
                unsafe_allow_html=True
            )

            ranking = gerar_ranking_bimestre(
                analise,
                bimestre_selecionado,
                disciplinas
            )

            if ranking.empty:

                st.info(
                    f"Ainda não há notas "
                    f"suficientes para o "
                    f"{bimestre_selecionado}."
                )

            else:

                ranking_exibicao = (
                    ranking.copy()
                )

                ranking_exibicao[
                    "Média do Bimestre"
                ] = (
                    ranking_exibicao[
                        "Média do Bimestre"
                    ]
                    .map(formatar_nota)
                )

                st.dataframe(
                    ranking_exibicao,
                    use_container_width=True,
                    hide_index=True
                )

        # ====================================================
        # ABA 3 - RECUPERAÇÕES
        # ====================================================

        with aba3:

            st.markdown(
                '<div class="titulo-secao">'
                'PREPARAÇÕES DE RECUPERAÇÃO'
                '</div>',
                unsafe_allow_html=True
            )

            rec_tab1, rec_tab2, rec_tab3 = st.tabs([

                "📝 1º SEMESTRE",

                "📝 2º SEMESTRE",

                "🚨 RECUPERAÇÃO FINAL"

            ])

            # ------------------------------------------------
            # 1º SEMESTRE
            # ------------------------------------------------

            with rec_tab1:

                st.markdown(
                    '<div class="subtitulo-secao">'
                    'Alunos em Recuperação — 1º Semestre'
                    '</div>',
                    unsafe_allow_html=True
                )

                rec1 = (
                    preparar_relatorio_recuperacao(
                        analise,
                        1
                    )
                )

                if rec1.empty:

                    st.success(
                        "Nenhum aluno em "
                        "recuperação no 1º semestre."
                    )

                else:

                    st.dataframe(
                        rec1,
                        use_container_width=True,
                        hide_index=True
                    )

                    st.divider()

                    st.markdown(
                        "### 📚 Recuperação por disciplina"
                    )

                    rec1_disc = (
                        gerar_recuperacao_por_disciplina(
                            analise,
                            1
                        )
                    )

                    st.dataframe(
                        rec1_disc,
                        use_container_width=True,
                        hide_index=True
                    )

                    st.markdown(
                        "### 👩‍🎓 Lista de alunos por disciplina"
                    )

                    lista_rec1 = (
                        gerar_lista_alunos_por_disciplina(
                            analise,
                            1
                        )
                    )

                    st.dataframe(
                        lista_rec1,
                        use_container_width=True,
                        hide_index=True
                    )

            # ------------------------------------------------
            # 2º SEMESTRE
            # ------------------------------------------------

            with rec_tab2:

                st.markdown(
                    '<div class="subtitulo-secao">'
                    'Alunos em Recuperação — 2º Semestre'
                    '</div>',
                    unsafe_allow_html=True
                )

                rec2 = (
                    preparar_relatorio_recuperacao(
                        analise,
                        2
                    )
                )

                if rec2.empty:

                    st.success(
                        "Nenhum aluno em "
                        "recuperação no 2º semestre."
                    )

                else:

                    st.dataframe(
                        rec2,
                        use_container_width=True,
                        hide_index=True
                    )

                    st.divider()

                    st.markdown(
                        "### 📚 Recuperação por disciplina"
                    )

                    rec2_disc = (
                        gerar_recuperacao_por_disciplina(
                            analise,
                            2
                        )
                    )

                    st.dataframe(
                        rec2_disc,
                        use_container_width=True,
                        hide_index=True
                    )

                    st.markdown(
                        "### 👩‍🎓 Lista de alunos por disciplina"
                    )

                    lista_rec2 = (
                        gerar_lista_alunos_por_disciplina(
                            analise,
                            2
                        )
                    )

                    st.dataframe(
                        lista_rec2,
                        use_container_width=True,
                        hide_index=True
                    )

            # ------------------------------------------------
            # RECUPERAÇÃO FINAL
            # ------------------------------------------------

            with rec_tab3:

                st.markdown(
                    '<div class="subtitulo-secao">'
                    'Preparação da Recuperação Final'
                    '</div>',
                    unsafe_allow_html=True
                )

                rec_final = (
                    preparar_recuperacao_final(
                        analise
                    )
                )

                if rec_final.empty:

                    st.success(
                        "Nenhum aluno em "
                        "recuperação final."
                    )

                else:

                    st.dataframe(
                        rec_final,
                        use_container_width=True,
                        hide_index=True
                    )
        # ====================================================
        # ABA 4 - ANÁLISE DETALHADA
        # ====================================================

        with aba4:

            st.markdown(
                '<div class="titulo-secao">'
                'ANÁLISE DETALHADA'
                '</div>',
                unsafe_allow_html=True
            )

            resultado_filtrado = mostrar_filtros(
                analise,
                "principal"
            )

            st.dataframe(
                resultado_filtrado,
                use_container_width=True,
                hide_index=True
            )


        # ====================================================
        # ABA 5 - RELATÓRIOS
        # ====================================================

        with aba5:

            st.markdown(
                '<div class="titulo-secao">'
                'EXPORTAÇÕES DE RELATÓRIOS'
                '</div>',
                unsafe_allow_html=True
            )

            rc1, rc2 = st.columns(2)

            with rc1:

                periodo_relatorio = st.selectbox(
                    "Período do relatório",
                    [
                        "1º Semestre",
                        "2º Semestre",
                        "Ano Letivo"
                    ],
                    key="periodo_relatorio"
                )

            with rc2:

                bimestre_relatorio = st.selectbox(
                    "Bimestre utilizado "
                    "para os alunos destaque",
                    BIMESTRES,
                    index=BIMESTRES.index(
                        st.session_state.bimestre_selecionado
                    ),
                    key="bimestre_relatorio"
                )

            ranking_relatorio = gerar_ranking_bimestre(
                analise,
                bimestre_relatorio,
                disciplinas
            )

            st.divider()

            # ====================================================
            # ADMINISTRADOR: WORD + PDF
            # ====================================================

            if usuario_eh_administrador():

                col_word, col_pdf = st.columns(2)

                # ------------------------------------------------
                # WORD
                # ------------------------------------------------

                with col_word:

                    st.markdown(
                        "### 📝 Relatório Word"
                    )

                    try:

                        relatorio_word = gerar_relatorio_word(
                            analise,
                            ranking_relatorio,
                            periodo_relatorio,
                            bimestre_relatorio
                        )

                        st.download_button(
                            label="📘 Baixar relatório em Word",
                            data=relatorio_word,
                            file_name=(
                                "RELATORIO_DESEMPENHO_"
                                "ACADEMICO.docx"
                            ),
                            mime=(
                                "application/vnd.openxmlformats-"
                                "officedocument.wordprocessingml.document"
                            ),
                            use_container_width=True,
                            key="download_word"
                        )

                    except Exception as erro_word:

                        st.error(
                            "Não foi possível "
                            "gerar o relatório Word."
                        )

                        st.exception(
                            erro_word
                        )

                # ------------------------------------------------
                # PDF
                # ------------------------------------------------

                with col_pdf:

                    st.markdown(
                        "### 📕 Relatório PDF"
                    )

                    try:

                        relatorio_pdf = gerar_relatorio_pdf(
                            analise,
                            ranking_relatorio,
                            periodo_relatorio,
                            bimestre_relatorio
                        )

                        st.download_button(
                            label="📕 Baixar relatório em PDF",
                            data=relatorio_pdf,
                            file_name=(
                                "RELATORIO_DESEMPENHO_"
                                "ACADEMICO.pdf"
                            ),
                            mime="application/pdf",
                            use_container_width=True,
                            key="download_pdf_admin"
                        )

                    except Exception as erro_pdf:

                        st.error(
                            "Não foi possível "
                            "gerar o relatório PDF."
                        )

                        st.exception(
                            erro_pdf
                        )

            else:

                # ====================================================
                # DEMAIS USUÁRIOS: SOMENTE PDF
                # ====================================================

                st.markdown(
                    "### 📕 Relatório PDF"
                )

                st.caption(
                    "Disponível para download conforme "
                    "as permissões do seu perfil."
                )

                try:

                    relatorio_pdf = gerar_relatorio_pdf(
                        analise,
                        ranking_relatorio,
                        periodo_relatorio,
                        bimestre_relatorio
                    )

                    st.download_button(
                        label="📕 Baixar relatório em PDF",
                        data=relatorio_pdf,
                        file_name=(
                            "RELATORIO_DESEMPENHO_"
                            "ACADEMICO.pdf"
                        ),
                        mime="application/pdf",
                        use_container_width=True,
                        key="download_pdf_usuario"
                    )

                except Exception as erro_pdf:

                    st.error(
                        "Não foi possível "
                        "gerar o relatório PDF."
                    )

                    st.exception(
                        erro_pdf
                    )


        # ====================================================
        # ABA 6 - EXPORTAR DADOS
        # ====================================================

        with aba6:

            # ====================================================
            # EXCLUSIVO DO ADMINISTRADOR
            # ====================================================

            if not usuario_eh_administrador():

                st.warning(
                    "🔒 Esta área é exclusiva para o perfil Administrador."
                )

                st.info(
                    "Para exportar relatórios, utilize a aba "
                    "de exportação em PDF."
                )

            else:

                st.markdown(
                    '<div class="titulo-secao">'
                    'EXPORTAÇÃO DOS DADOS ANALISADOS'
                    '</div>',
                    unsafe_allow_html=True
                )

                st.write(
                    "Baixe os dados completos da análise "
                    "em formato Excel."
                )

                buffer_excel = io.BytesIO()

                try:

                    # ====================================================
                    # CRIAÇÃO DO ARQUIVO EXCEL
                    # ====================================================

                    with pd.ExcelWriter(
                        buffer_excel,
                        engine="openpyxl"
                    ) as writer:

                        # ============================================
                        # ANÁLISE COMPLETA
                        # ============================================

                        analise.to_excel(
                            writer,
                            sheet_name="Análise",
                            index=False
                        )

                        # ============================================
                        # RESUMOS
                        # ============================================

                        resumo_alunos.to_excel(
                            writer,
                            sheet_name="Resumo Alunos",
                            index=False
                        )

                        resumo_disciplinas.to_excel(
                            writer,
                            sheet_name="Resumo Disciplinas",
                            index=False
                        )

                        # ============================================
                        # RECUPERAÇÕES
                        # ============================================

                        rec1 = preparar_relatorio_recuperacao(
                            analise,
                            1
                        )

                        rec2 = preparar_relatorio_recuperacao(
                            analise,
                            2
                        )

                        rec_final = preparar_recuperacao_final(
                            analise
                        )

                        rec1.to_excel(
                            writer,
                            sheet_name="Recuperação 1 Sem",
                            index=False
                        )

                        rec2.to_excel(
                            writer,
                            sheet_name="Recuperação 2 Sem",
                            index=False
                        )

                        rec_final.to_excel(
                            writer,
                            sheet_name="Recuperação Final",
                            index=False
                        )

                        # ============================================
                        # RANKINGS DOS BIMESTRES
                        # ============================================

                        for bimestre in BIMESTRES:

                            ranking = gerar_ranking_bimestre(
                                analise,
                                bimestre,
                                disciplinas
                            )

                            nome_aba = (
                                bimestre
                                .replace("º", "")
                                .replace(" ", "_")
                            )

                            ranking.to_excel(
                                writer,
                                sheet_name=f"Top5_{nome_aba}"[:31],
                                index=False
                            )

                    buffer_excel.seek(0)

                    # ============================================
                    # DOWNLOAD DO EXCEL
                    # ============================================

                    st.download_button(
                        label="📊 Baixar análise completa em Excel",
                        data=buffer_excel.getvalue(),
                        file_name=(
                            "ANALISE_COMPLETA_NOTAS_SESI.xlsx"
                        ),
                        mime=(
                            "application/vnd.openxmlformats-officedocument"
                            ".spreadsheetml.sheet"
                        ),
                        use_container_width=True,
                        key="download_excel_completo"
                    )

                    st.divider()

                    st.info(
                        "O arquivo Excel contém as abas "
                        "de análise, resumos, recuperações "
                        "e rankings dos quatro bimestres."
                    )

                except Exception as erro_excel:

                    st.error(
                        "Não foi possível gerar o arquivo Excel."
                    )

                    st.exception(
                        erro_excel
                    )


    # ============================================================
    # ERRO GERAL DO PROCESSAMENTO
    # ============================================================

    except Exception as erro:

        st.error(
            "❌ Erro ao executar as análises."
        )

        st.exception(
            erro
        )